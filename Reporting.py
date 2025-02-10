import streamlit as st
import fitz  # PyMuPDF for PDF handling
import docx
import re
import mammoth
import pdfkit
from io import BytesIO

def convert_docx_to_pdf(docx_path, pdf_output_path):
    """Convert a Word document to PDF using Mammoth and pdfkit."""
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html_content = result.value

    html_path = docx_path.replace(".docx", ".html")
    pdf_path = docx_path.replace(".docx", ".pdf")

    with open(html_path, "w") as html_file:
        html_file.write(html_content)

    # Specify the path for wkhtmltopdf
    config = pdfkit.configuration(wkhtmltopdf="/usr/bin/wkhtmltopdf")

    pdfkit.from_file(html_path, pdf_path, configuration=config)
    return pdf_path

def extract_placeholders(doc_path):
    """Extract placeholders from a Word document."""
    doc = docx.Document(doc_path)
    placeholders = set()
    pattern = re.compile(r'\{(.*?)\}')  # Match {placeholder}
    
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        placeholders.update(matches)
    
    st.write("**Detected Placeholders:**", placeholders)  # Debugging output in Streamlit
    return list(placeholders)

def replace_placeholders(doc_path, values):
    """Replace placeholders in a Word document."""
    doc = docx.Document(doc_path)
    for para in doc.paragraphs:
        for key, val in values.items():
            if f'{{{key}}}' in para.text:
                para.text = para.text.replace(f'{{{key}}}', val)
    
    output_path = "filled_template.docx"
    doc.save(output_path)
    return output_path

def merge_pdfs(main_pdf, annex_pdfs, output_pdf):
    """Merge the main report PDF with annex PDFs."""
    merger = fitz.open()
    merger.insert_pdf(fitz.open(main_pdf))
    for annex in annex_pdfs:
        merger.insert_pdf(fitz.open(annex))
    merger.save(output_pdf)
    merger.close()
    return output_pdf

# Streamlit UI
st.title("Caisson Report Generator")
st.write("Fill in the details and upload annexes to generate the final report.")

# Upload template
template_file = st.file_uploader("Upload Word template (DOCX)", type=["docx"])

values = {}
if template_file:
    docx_path = "temp_template.docx"
    with open(docx_path, "wb") as f:
        f.write(template_file.read())
    
    placeholders = extract_placeholders(docx_path)
    if placeholders:
        st.sidebar.header("Enter Report Details")
        for ph in placeholders:
            values[ph] = st.sidebar.text_input(f"{ph}")
    else:
        st.warning("No placeholders found in the uploaded document. Make sure placeholders are enclosed in {}.")
    
    # Upload annex PDFs
    annex_files = st.file_uploader("Upload Annex PDFs", type=["pdf"], accept_multiple_files=True)
    
    # Generate button
    if st.button("Generate Report") and values:
        filled_docx = replace_placeholders(docx_path, values)
        pdf_report = "final_report.pdf"
        pdf_report = convert_docx_to_pdf(filled_docx, pdf_report)  # Ensure variable is updated
        
        annex_paths = []
        for annex in annex_files:
            annex_path = f"annex_{annex.name}"
            with open(annex_path, "wb") as f:
                f.write(annex.read())
            annex_paths.append(annex_path)
        
        final_pdf = "complete_report.pdf"
        merge_pdfs(pdf_report, annex_paths, final_pdf)
        
        st.success("Report Generated Successfully!")
        with open(final_pdf, "rb") as f:
            st.download_button("Download Final Report", f, file_name="Caisson_Report.pdf")
